import subprocess
import sys
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from docx import Document
from src.common import ENV_KEYS, config_output_path
from src.docx import build_docx
from src.html import build_html
from src.md.render_md import render_markdown
from src.pdf import build_pdf
from tests.conftest import DEFAULT_CONFIG, DUMMY_EMAIL, DUMMY_NAME, SAMPLE_YAML

SCRIPT_DIR = Path(__file__).resolve().parent.parent


def _add_run(paragraph, text):
    run = paragraph.add_run(text)
    return run


# ─── _format_name_header ──────────────────────────────────────


def test_name_header_title_style():
    doc = Document()
    p = doc.add_paragraph(DUMMY_NAME)
    p.style = doc.styles["Title"]
    _add_run(p, DUMMY_NAME)
    assert build_docx._format_name_header(p, DUMMY_NAME)


def test_name_header_heading1_style():
    doc = Document()
    p = doc.add_paragraph(f"{DUMMY_NAME} - CEO")
    p.style = doc.styles["Heading 1"]
    _add_run(p, f"{DUMMY_NAME} - CEO")
    assert build_docx._format_name_header(p, DUMMY_NAME)


def test_name_header_no_name():
    doc = Document()
    p = doc.add_paragraph("Not My Name")
    p.style = doc.styles["Title"]
    _add_run(p, "Not My Name")
    assert not build_docx._format_name_header(p, DUMMY_NAME)


def test_name_header_wrong_style_and_long():
    doc = Document()
    p = doc.add_paragraph(f"{DUMMY_NAME} - Chief Executive Officer")
    p.style = doc.styles["Heading 2"]
    _add_run(p, f"{DUMMY_NAME} - Chief Executive Officer")
    assert not build_docx._format_name_header(p, DUMMY_NAME)


def test_name_header_wrong_style_but_short():
    doc = Document()
    p = doc.add_paragraph(DUMMY_NAME)
    p.style = doc.styles["Normal"]
    assert build_docx._format_name_header(p, DUMMY_NAME)


def test_name_header_formats_correctly():
    doc = Document()
    p = doc.add_paragraph(DUMMY_NAME)
    p.style = doc.styles["Title"]
    _add_run(p, DUMMY_NAME)
    assert build_docx._format_name_header(p, DUMMY_NAME)
    assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert p.runs[0].font.size == Pt(26)
    assert p.runs[0].font.bold
    assert p.runs[0].font.all_caps


# ─── _format_contact_info ─────────────────────────────────────


def test_contact_info_styling():
    doc = Document()
    p = doc.add_paragraph("test@example.com")
    _add_run(p, "test@example.com")
    build_docx._format_contact_info(p)
    assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert p.paragraph_format.space_after == Pt(14)
    assert p.runs[0].font.color.rgb == RGBColor(0x55, 0x55, 0x55)


# ─── _format_heading2 ─────────────────────────────────────────


def test_heading2_styling():
    doc = Document()
    p = doc.add_paragraph("Experience")
    _add_run(p, "Experience")
    build_docx._format_heading2(p)
    assert p.runs[0].font.size == Pt(13)
    assert p.runs[0].font.color.rgb == RGBColor(0x29, 0x80, 0xB9)
    assert p.runs[0].font.all_caps
    assert p.runs[0].font.bold
    assert p.paragraph_format.space_before == Pt(16)
    assert p.paragraph_format.space_after == Pt(5)


# ─── _format_heading3 ─────────────────────────────────────────


def test_heading3_styling():
    doc = Document()
    p = doc.add_paragraph("Company Name")
    _add_run(p, "Company Name")
    build_docx._format_heading3(p)
    assert p.runs[0].font.size == Pt(11)
    assert p.runs[0].font.color.rgb == RGBColor(0x1A, 0x25, 0x2F)
    assert p.runs[0].font.bold
    assert p.paragraph_format.space_before == Pt(10)
    assert p.paragraph_format.space_after == Pt(2)


def test_heading3_contacts_page_break():
    doc = Document()
    p = doc.add_paragraph("Contacts+")
    _add_run(p, "Contacts+")
    build_docx._format_heading3(p)
    assert p.paragraph_format.page_break_before


def test_heading3_no_page_break():
    doc = Document()
    p = doc.add_paragraph("Other Company")
    _add_run(p, "Other Company")
    build_docx._format_heading3(p)
    assert not p.paragraph_format.page_break_before


# ─── _format_job_title ────────────────────────────────────────


def test_job_title_styling():
    doc = Document()
    p = doc.add_paragraph("Engineer | 2020-2025")
    _add_run(p, "Engineer | 2020-2025")
    build_docx._format_job_title(p)
    assert p.paragraph_format.space_before == Pt(6)
    assert p.paragraph_format.space_after == Pt(1)


# ─── helpers ────────────────────────────────────────────────


_REVERSE_KEYS = {v: k for k, v in ENV_KEYS.items()}


def _write_env(path, data):
    lines = [f"{_REVERSE_KEYS[k]}={v}" for k, v in data.items() if k in _REVERSE_KEYS]
    path.write_text("\n".join(lines) + "\n")


def _write_minimal_yaml(path):
    path.write_text(
        "basics:\n"
        "  name: '{{ name }}'\n"
        "  email: '{{ email }}'\n"
        "  summary: Summary.\n"
        "  location:\n"
        "    city: City\n"
        "    region: Region\n"
        "work: []\n"
        "early_career: []\n"
        "skills: []\n"
        "education: []"
    )


def _write_minimal_yaml_public(path):
    path.write_text(
        "basics:\n"
        "  name: '{{ name }}'\n"
        "  email: '{{ email }}'\n"
        "  phone: '{{ phone }}'\n"
        "  summary: Summary.\n"
        "  location:\n"
        "    city: City\n"
        "    region: Region\n"
        "work: []\n"
        "early_career: []\n"
        "skills: []\n"
        "education: []"
    )


def test_build_flawless_pdf_happy_path(tmp_path):
    _write_env(tmp_path / ".env.local", DEFAULT_CONFIG)
    yaml_file = tmp_path / "test.yaml"
    yaml_file.write_text(SAMPLE_YAML)
    build_pdf.build_flawless_pdf(str(yaml_file))
    pdf_file = tmp_path / "john_doe_resume.pdf"
    assert pdf_file.exists()
    assert pdf_file.stat().st_size > 100
    assert pdf_file.read_bytes()[:4] == b"%PDF"


def test_build_flawless_pdf_missing_file(tmp_path, capsys):
    missing = tmp_path / "nonexistent.yaml"
    build_pdf.build_flawless_pdf(str(missing))
    captured = capsys.readouterr()
    assert "not found" in captured.out


def test_build_flawless_pdf_with_config(tmp_path):
    _write_env(tmp_path / ".env.local", {"name": "Jane Doe", "email": "j@e.co"})
    yaml_file = tmp_path / "cv.yaml"
    _write_minimal_yaml(yaml_file)
    build_pdf.build_flawless_pdf(str(yaml_file))
    assert (tmp_path / "jane_doe_resume.pdf").exists()


def test_build_flawless_pdf_public(tmp_path):
    cfg = {"name": "Jane Doe", "email": "j@e.co", "phone": "555-0000"}
    _write_env(tmp_path / ".env.local", cfg)
    yaml_file = tmp_path / "cv.yaml"
    _write_minimal_yaml_public(yaml_file)
    build_pdf.build_flawless_pdf(str(yaml_file), public=True)
    assert (tmp_path / "jane_doe_resume_public.pdf").exists()
    size_public = (tmp_path / "jane_doe_resume_public.pdf").stat().st_size
    assert size_public > 100


# ─── build_styled_docx happy path ─────────────────────────────


def test_build_styled_docx_happy_path(tmp_path):
    _write_env(tmp_path / ".env.local", DEFAULT_CONFIG)
    yaml_file = tmp_path / "test.yaml"
    yaml_file.write_text(SAMPLE_YAML)
    build_docx.build_styled_docx(str(yaml_file))
    docx_file = tmp_path / "john_doe_resume.docx"
    assert docx_file.exists()
    assert docx_file.stat().st_size > 100
    assert docx_file.read_bytes()[:2] == b"PK"


def test_build_styled_docx_missing_file(tmp_path, capsys):
    missing = tmp_path / "nonexistent.yaml"
    build_docx.build_styled_docx(str(missing))
    captured = capsys.readouterr()
    assert "not found" in captured.out


def test_build_styled_docx_with_config(tmp_path):
    _write_env(tmp_path / ".env.local", {"name": "Jane Doe", "email": "j@e.co"})
    yaml_file = tmp_path / "cv.yaml"
    _write_minimal_yaml(yaml_file)
    build_docx.build_styled_docx(str(yaml_file))
    assert (tmp_path / "jane_doe_resume.docx").exists()


# ─── build_web_html happy path ────────────────────────────────


def test_build_html_happy_path(tmp_path):
    _write_env(tmp_path / ".env.local", DEFAULT_CONFIG)
    yaml_file = tmp_path / "test.yaml"
    yaml_file.write_text(SAMPLE_YAML)
    build_html.build_web_html(str(yaml_file))
    html_file = tmp_path / "john_doe_resume.html"
    assert html_file.exists()
    assert html_file.stat().st_size > 100
    content = html_file.read_text()
    assert "<!DOCTYPE html>" in content
    assert DUMMY_NAME in content
    assert f'href="mailto:{DUMMY_EMAIL}"' in content
    assert "john_doe_resume.pdf" in content


def test_build_html_missing_file(tmp_path, capsys):
    missing = tmp_path / "nonexistent.yaml"
    build_html.build_web_html(str(missing))
    captured = capsys.readouterr()
    assert "not found" in captured.out


def test_build_html_with_config(tmp_path):
    _write_env(tmp_path / ".env.local", {"name": "Jane Doe", "email": "j@e.co"})
    yaml_file = tmp_path / "cv.yaml"
    _write_minimal_yaml(yaml_file)
    build_html.build_web_html(str(yaml_file))
    html_file = tmp_path / "jane_doe_resume.html"
    assert html_file.exists()
    content = html_file.read_text()
    assert "Jane Doe" in content


# ─── CLI entry points ─────────────────────────────────────────


def test_cli_no_args_pdf():
    result = subprocess.run(
        [sys.executable, "src/pdf/build_pdf.py"],
        capture_output=True,
        text=True,
        cwd=SCRIPT_DIR,
    )
    assert result.returncode == 1
    assert "Usage:" in result.stdout


def test_cli_no_args_docx():
    result = subprocess.run(
        [sys.executable, "src/docx/build_docx.py"],
        capture_output=True,
        text=True,
        cwd=SCRIPT_DIR,
    )
    assert result.returncode == 1
    assert "Usage:" in result.stdout


def test_cli_no_args_html():
    result = subprocess.run(
        [sys.executable, "src/html/build_html.py"],
        capture_output=True,
        text=True,
        cwd=SCRIPT_DIR,
    )
    assert result.returncode == 1
    assert "Usage:" in result.stdout


def test_config_output_path_with_name_no_output_dir():
    got = config_output_path("/tmp/cv.md", {"name": "Jane Doe"}, "pdf")
    assert got == "/tmp/jane_doe_resume.pdf"


def test_config_output_path_with_name_and_output_dir():
    got = config_output_path("/tmp/cv.md", {"name": "Jane Doe"}, "pdf", output_dir="/out")
    assert got == "/out/jane_doe_resume.pdf"


def test_config_output_path_no_name_fallback():
    got = config_output_path("/tmp/cv.md", {"email": "j@e.co"}, "pdf")
    assert got == "/tmp/cv.pdf"


def test_render_markdown(tmp_path):
    _write_env(tmp_path / ".env.local", DEFAULT_CONFIG)
    yaml_file = tmp_path / "test.yaml"
    yaml_file.write_text(SAMPLE_YAML)
    output_md = tmp_path / "output.md"
    render_markdown(str(yaml_file), str(output_md))
    assert output_md.exists()
    content = output_md.read_text()
    assert DUMMY_NAME in content
    assert "Company Name" in content


def test_cli_no_args_md():
    result = subprocess.run(
        [sys.executable, "src/md/render_md.py"],
        capture_output=True,
        text=True,
        cwd=SCRIPT_DIR,
    )
    assert result.returncode == 1
    assert "Usage:" in result.stdout
